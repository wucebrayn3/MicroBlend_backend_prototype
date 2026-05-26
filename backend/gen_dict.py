from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ---- Title ----
title = doc.add_heading("MicroBlend — Database Data Dictionary", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("System version: 1.0")
doc.add_paragraph("")

# ---- Data ----
# Each entry: (table, field, data_type, description, explanation)
entries = [
    # USERS
    ("users", "id", "BIGINT (PK)", "Primary key, auto-increment", "Unique identifier for every user record"),
    ("users", "username", "VARCHAR(150)", "Login username, unique", "Used for authentication; Django AbstractUser requirement"),
    ("users", "password", "VARCHAR(128)", "Hashed password", "Stored as Django PBKDF2 hash; never plaintext"),
    ("users", "email", "VARCHAR(254)", "Email address, unique, nullable", "Used for contact and password reset; blank allowed for guests"),
    ("users", "phone", "VARCHAR(20)", "Phone number, unique, nullable", "Contact number; alternative identifier for customer lookup"),
    ("users", "role", "VARCHAR(20)", "User role: admin/staff/customer", "Determines permissions and UI access level"),
    ("users", "staff_role", "VARCHAR(20)", "Staff subtype: waiter/cashier/kitchen/bar, nullable", "Further refines staff permissions; null for non-staff"),
    ("users", "registered_device_id", "VARCHAR(128)", "Registered device identifier, unique, nullable", "Links a guest to their device for session continuity"),
    ("users", "is_guest", "BOOLEAN", "Guest account flag", "True for walk-in guests who did not create a full account"),
    ("users", "guest_expires_at", "DATETIME", "Guest account expiration, nullable", "Auto-cleaned after expiry to prevent orphaned guest records"),
    ("users", "is_active", "BOOLEAN", "Account active flag", "False disables login without deleting the record"),
    ("users", "is_deleted", "BOOLEAN", "Soft-delete flag", "True hides the user from queries; preserves FK integrity"),
    ("users", "is_staff", "BOOLEAN", "Django admin access flag", "Grants access to Django admin site"),
    ("users", "is_superuser", "BOOLEAN", "Superuser flag", "Grants all permissions without explicit assignment"),
    ("users", "first_name", "VARCHAR(150)", "Given name", "Display name component; Django built-in field"),
    ("users", "last_name", "VARCHAR(150)", "Surname", "Display name component; Django built-in field"),
    ("users", "date_joined", "DATETIME", "Account creation timestamp", "Set automatically on user creation"),
    ("users", "last_login", "DATETIME", "Last successful login timestamp, nullable", "Updated on each authentication"),
    ("users", "created_at", "DATETIME", "Record creation timestamp", "Inherited from BaseModel; set on INSERT"),
    ("users", "updated_at", "DATETIME", "Last modification timestamp", "Inherited from BaseModel; updated on every UPDATE"),

    # CATEGORIES
    ("categories", "id", "BIGINT (PK)", "Primary key, auto-increment", "Unique identifier"),
    ("categories", "name", "VARCHAR(100)", "Category name, unique", "Display name for menu category groupings"),
    ("categories", "sort_order", "INTEGER", "Display sort order", "Controls ordering in menu listings; lower = first"),
    ("categories", "created_at", "DATETIME", "Record creation timestamp", "Inherited from BaseModel"),
    ("categories", "updated_at", "DATETIME", "Last modification timestamp", "Inherited from BaseModel"),

    # MENU ITEMS
    ("menu_items", "id", "BIGINT (PK)", "Primary key, auto-increment", "Unique identifier"),
    ("menu_items", "name", "VARCHAR(200)", "Item name, unique", "Display name on menus and receipts"),
    ("menu_items", "description", "TEXT", "Item description, blankable", "Detailed item description shown on digital menus"),
    ("menu_items", "price", "DECIMAL(10,2)", "Selling price", "Current price of the menu item"),
    ("menu_items", "preparation_station", "VARCHAR(20)", "Station: kitchen/bar/none", "Determines which station fulfills this item"),
    ("menu_items", "prep_eta_minutes", "INTEGER", "Estimated preparation time in minutes", "Shown to customers as wait time estimate"),
    ("menu_items", "is_available", "BOOLEAN", "Available for ordering flag", "False when ingredient stock is insufficient or item is paused"),
    ("menu_items", "popularity_score", "INTEGER", "Popularity ranking counter", "Incremented on order; used for sorting and analytics"),
    ("menu_items", "category_id", "BIGINT (FK)", "FK to categories.id, nullable", "SET_NULL on category delete; item survives category removal"),
    ("menu_items", "created_at", "DATETIME", "Record creation timestamp", "Inherited from BaseModel"),
    ("menu_items", "updated_at", "DATETIME", "Last modification timestamp", "Inherited from BaseModel"),

    # MENU ITEM INGREDIENTS (junction)
    ("menu_item_ingredients", "id", "BIGINT (PK)", "Primary key, auto-increment", "Unique identifier"),
    ("menu_item_ingredients", "menu_item_id", "BIGINT (FK)", "FK to menu_items.id", "CASCADE on menu item delete; ingredient link is removed with item"),
    ("menu_item_ingredients", "ingredient_id", "BIGINT (FK)", "FK to ingredients.id", "CASCADE on ingredient delete; item becomes orphaned if ingredient removed"),
    ("menu_item_ingredients", "quantity_required", "DECIMAL(10,2)", "Amount of ingredient needed", "Used for inventory deduction calculations when order is placed"),
    ("menu_item_ingredients", "created_at", "DATETIME", "Record creation timestamp", "Inherited from BaseModel"),
    ("menu_item_ingredients", "updated_at", "DATETIME", "Last modification timestamp", "Inherited from BaseModel"),

    # INGREDIENTS
    ("ingredients", "id", "BIGINT (PK)", "Primary key, auto-increment", "Unique identifier"),
    ("ingredients", "name", "VARCHAR(150)", "Ingredient name, unique", "Display name for inventory tracking"),
    ("ingredients", "unit", "VARCHAR(10)", "Unit of measure: g/ml/pc", "Determines how quantities are measured and displayed"),
    ("ingredients", "reorder_level", "DECIMAL(10,2)", "Low-stock threshold", "When quantity_remaining drops below this, a reorder alert triggers"),
    ("ingredients", "created_at", "DATETIME", "Record creation timestamp", "Inherited from BaseModel"),
    ("ingredients", "updated_at", "DATETIME", "Last modification timestamp", "Inherited from BaseModel"),

    # INVENTORY BATCHES
    ("inventory_batches", "id", "BIGINT (PK)", "Primary key, auto-increment", "Unique identifier"),
    ("inventory_batches", "ingredient_id", "BIGINT (FK)", "FK to ingredients.id", "CASCADE; batch is removed if ingredient is deleted"),
    ("inventory_batches", "quantity_added", "DECIMAL(10,2)", "Initial quantity received", "Used to track batch yield and waste percentage"),
    ("inventory_batches", "quantity_remaining", "DECIMAL(10,2)", "Current quantity left", "Decremented as orders consume this batch; FIFO-based deduction"),
    ("inventory_batches", "unit_cost", "DECIMAL(10,2)", "Cost per unit", "Used for COGS (Cost of Goods Sold) calculations"),
    ("inventory_batches", "expiration_date", "DATE", "Batch expiry date", "Alerts staff when approaching expiry; used for FIFO rotation"),
    ("inventory_batches", "source", "VARCHAR(120)", "Supplier or delivery note, nullable", "Audit trail for batch origin; useful for recall scenarios"),
    ("inventory_batches", "created_at", "DATETIME", "Record creation timestamp", "Inherited from BaseModel"),
    ("inventory_batches", "updated_at", "DATETIME", "Last modification timestamp", "Inherited from BaseModel"),

    # INVENTORY MOVEMENTS
    ("inventory_movements", "id", "BIGINT (PK)", "Primary key, auto-increment", "Unique identifier"),
    ("inventory_movements", "ingredient_id", "BIGINT (FK)", "FK to ingredients.id", "CASCADE; movement log removed with ingredient"),
    ("inventory_movements", "batch_id", "BIGINT (FK)", "FK to inventory_batches.id, nullable", "SET_NULL; movement survives batch deletion"),
    ("inventory_movements", "movement_type", "VARCHAR(20)", "Type: restock/deduct/adjust/delete", "Categorizes the inventory transaction"),
    ("inventory_movements", "quantity", "DECIMAL(10,2)", "Quantity moved", "Positive for additions, negative for deductions"),
    ("inventory_movements", "actor_id", "BIGINT (FK)", "FK to users.id, nullable", "SET_NULL; identifies who performed the movement"),
    ("inventory_movements", "note", "VARCHAR(255)", "Optional note, nullable", "Free-text reason for the movement"),
    ("inventory_movements", "related_order_id", "BIGINT", "Denormalized order reference, nullable", "Links movement back to the order that caused it; not a FK for performance"),
    ("inventory_movements", "created_at", "DATETIME", "Record creation timestamp", "Inherited from BaseModel"),
    ("inventory_movements", "updated_at", "DATETIME", "Last modification timestamp", "Inherited from BaseModel"),

    # ORDERS
    ("orders", "id", "BIGINT (PK)", "Primary key, auto-increment", "Unique identifier"),
    ("orders", "table_session_id", "BIGINT (FK)", "FK to table_sessions.id, nullable", "SET_NULL; order survives session close"),
    ("orders", "placed_by_id", "BIGINT (FK)", "FK to users.id, nullable", "SET_NULL; order survives user deletion"),
    ("orders", "placed_by_name", "VARCHAR(120)", "Denormalized name of who placed it, nullable", "Display name even if user is later deleted"),
    ("orders", "placed_by_role", "VARCHAR(30)", "Role at time of placement, nullable", "Audit context: who placed the order"),
    ("orders", "placed_for_name", "VARCHAR(120)", "Customer name this order is for, nullable", "Used when staff places on behalf of a customer"),
    ("orders", "placed_for_contact", "VARCHAR(120)", "Customer contact info, nullable", "Used for delivery or customer lookup"),
    ("orders", "channel", "VARCHAR(30)", "Order channel: customer_account/guest/waiter_assisted/bulk/pos_sync", "Identifies how the order entered the system"),
    ("orders", "status", "VARCHAR(30)", "Overall order status: draft/placed/preparing/ready/payment_pending/paid/cancelled", "Top-level order lifecycle state"),
    ("orders", "kitchen_status", "VARCHAR(30)", "Kitchen fulfillment status", "Parallel track; independent of bar and cashier status"),
    ("orders", "bar_status", "VARCHAR(30)", "Bar fulfillment status", "Parallel track; for drink items"),
    ("orders", "cashier_status", "VARCHAR(30)", "Cashier fulfillment status", "Parallel track; for payment lifecycle"),
    ("orders", "is_bulk_order", "BOOLEAN", "Bulk order flag", "True for large pre-orders; requires cashier verification"),
    ("orders", "requires_cashier_verification", "BOOLEAN", "Cashier verification needed flag", "True for high-value or bulk orders; blocks completion until approved"),
    ("orders", "price_alert_sent", "BOOLEAN", "Price change alert sent flag", "Prevents duplicate alerts when menu prices change on active orders"),
    ("orders", "total_amount", "DECIMAL(10,2)", "Order total", "Sum of all line items; recomputed when items change"),
    ("orders", "receipt_number", "VARCHAR(40)", "Unique receipt number, nullable", "Auto-generated; customer-facing order reference"),
    ("orders", "notes", "TEXT", "Order-level notes, blankable", "Special instructions for the entire order"),
    ("orders", "external_pos_reference", "VARCHAR(120)", "Reference from external POS system, nullable", "Maps this order to the source system's order ID"),
    ("orders", "inventory_committed", "BOOLEAN", "Inventory deduction flag", "True once ingredients are deducted; prevents double-deduction"),
    ("orders", "created_at", "DATETIME", "Record creation timestamp", "Inherited from BaseModel"),
    ("orders", "updated_at", "DATETIME", "Last modification timestamp", "Inherited from BaseModel"),

    # ORDER ITEMS
    ("order_items", "id", "BIGINT (PK)", "Primary key, auto-increment", "Unique identifier"),
    ("order_items", "order_id", "BIGINT (FK)", "FK to orders.id", "CASCADE; items removed with order"),
    ("order_items", "menu_item_id", "BIGINT (FK)", "FK to menu_items.id, nullable", "SET_NULL; historical order items survive menu changes"),
    ("order_items", "item_name", "VARCHAR(200)", "Denormalized item name at time of order", "Snapshot so receipts show original name even if menu item is renamed"),
    ("order_items", "station", "VARCHAR(20)", "Station: kitchen/bar/none", "Copied from menu_item at order time; routes item to correct station"),
    ("order_items", "quantity", "INTEGER", "Quantity ordered", "Number of this item in the order"),
    ("order_items", "unit_price", "DECIMAL(10,2)", "Price at time of order", "Snapshot so price changes don't affect existing orders"),
    ("order_items", "customization_notes", "VARCHAR(255)", "Customization instructions, nullable", "Customer special requests for this specific item"),
    ("order_items", "created_at", "DATETIME", "Record creation timestamp", "Inherited from BaseModel"),
    ("order_items", "updated_at", "DATETIME", "Last modification timestamp", "Inherited from BaseModel"),

    # ORDER STATUS LOGS
    ("order_status_logs", "id", "BIGINT (PK)", "Primary key, auto-increment", "Unique identifier"),
    ("order_status_logs", "order_id", "BIGINT (FK)", "FK to orders.id", "CASCADE; log removed with order"),
    ("order_status_logs", "actor_id", "BIGINT (FK)", "FK to users.id, nullable", "SET_NULL; log survives user deletion"),
    ("order_status_logs", "status", "VARCHAR(40)", "Status value at time of transition", "Immutable record of what the status changed to"),
    ("order_status_logs", "note", "VARCHAR(255)", "Transition note, nullable", "Reason or context for the status change"),
    ("order_status_logs", "metadata", "JSON", "Arbitrary JSON payload", "Flexible extra data: timestamps, station info, etc."),
    ("order_status_logs", "created_at", "DATETIME", "Record creation timestamp", "Inherited from BaseModel"),
    ("order_status_logs", "updated_at", "DATETIME", "Last modification timestamp", "Inherited from BaseModel"),

    # FEEDBACK ENTRIES
    ("feedback_entries", "id", "BIGINT (PK)", "Primary key, auto-increment", "Unique identifier"),
    ("feedback_entries", "submitted_by_id", "BIGINT (FK)", "FK to users.id, nullable", "SET_NULL; feedback survives user deletion"),
    ("feedback_entries", "order_id", "BIGINT (FK)", "FK to orders.id, nullable", "SET_NULL; feedback survives order deletion"),
    ("feedback_entries", "entry_type", "VARCHAR(20)", "Type: feedback/report", "Distinguishes general feedback from problem reports"),
    ("feedback_entries", "subject", "VARCHAR(150)", "Short subject line", "Brief summary of the feedback"),
    ("feedback_entries", "message", "TEXT", "Full feedback message", "Detailed text from the customer"),
    ("feedback_entries", "status", "VARCHAR(20)", "Status: open/reviewed/resolved", "Tracks handling lifecycle of the feedback"),
    ("feedback_entries", "created_at", "DATETIME", "Record creation timestamp", "Inherited from BaseModel"),
    ("feedback_entries", "updated_at", "DATETIME", "Last modification timestamp", "Inherited from BaseModel"),

    # ORDER PLAYLISTS
    ("order_playlists", "id", "BIGINT (PK)", "Primary key, auto-increment", "Unique identifier"),
    ("order_playlists", "owner_id", "BIGINT (FK)", "FK to users.id", "CASCADE; playlist removed if user is deleted"),
    ("order_playlists", "name", "VARCHAR(120)", "Playlist name", "Customer-facing name for a saved order template"),
    ("order_playlists", "created_at", "DATETIME", "Record creation timestamp", "Inherited from BaseModel"),
    ("order_playlists", "updated_at", "DATETIME", "Last modification timestamp", "Inherited from BaseModel"),

    # ORDER PLAYLIST ITEMS
    ("order_playlist_items", "id", "BIGINT (PK)", "Primary key, auto-increment", "Unique identifier"),
    ("order_playlist_items", "playlist_id", "BIGINT (FK)", "FK to order_playlists.id", "CASCADE; item removed with playlist"),
    ("order_playlist_items", "menu_item_id", "BIGINT (FK)", "FK to menu_items.id", "CASCADE; playlist item removed if menu item is deleted"),
    ("order_playlist_items", "quantity", "INTEGER", "Default quantity in playlist", "How many of this item when the playlist is ordered"),
    ("order_playlist_items", "created_at", "DATETIME", "Record creation timestamp", "Inherited from BaseModel"),
    ("order_playlist_items", "updated_at", "DATETIME", "Last modification timestamp", "Inherited from BaseModel"),

    # TABLES
    ("tables", "id", "BIGINT (PK)", "Primary key, auto-increment", "Unique identifier"),
    ("tables", "identifier", "VARCHAR(30)", "Table display name/number, unique", "Customer-facing table identifier (e.g. 'Table 5')"),
    ("tables", "capacity", "INTEGER", "Maximum seating capacity", "Used for party-size validation when seating"),
    ("tables", "status", "VARCHAR(20)", "Status: vacant/occupied/merged/blocked", "Current table state for the floor plan"),
    ("tables", "zone", "VARCHAR(100)", "Section/zone name, nullable", "Groups tables by area (e.g. 'Patio', 'Bar')"),
    ("tables", "qr_code_value", "VARCHAR(64)", "Unique QR code value", "Auto-generated; embedded in QR code for customer scanning"),
    ("tables", "created_at", "DATETIME", "Record creation timestamp", "Inherited from BaseModel"),
    ("tables", "updated_at", "DATETIME", "Last modification timestamp", "Inherited from BaseModel"),

    # TABLE MERGE GROUPS
    ("table_merge_groups", "id", "BIGINT (PK)", "Primary key, auto-increment", "Unique identifier"),
    ("table_merge_groups", "name", "VARCHAR(64)", "Group name, unique", "Auto-generated label for merged table groups"),
    ("table_merge_groups", "is_active", "BOOLEAN", "Group active flag", "False when merged tables are split apart"),
    ("table_merge_groups", "created_at", "DATETIME", "Record creation timestamp", "Inherited from BaseModel"),
    ("table_merge_groups", "updated_at", "DATETIME", "Last modification timestamp", "Inherited from BaseModel"),

    # TABLE MERGE GROUP TABLES (junction)
    ("table_merge_group_tables", "id", "BIGINT (PK)", "Primary key, auto-increment", "Unique identifier"),
    ("table_merge_group_tables", "table_merge_group_id", "BIGINT (FK)", "FK to table_merge_groups.id", "Links a table to a merge group"),
    ("table_merge_group_tables", "table_id", "BIGINT (FK)", "FK to tables.id", "Links a merge group to a table"),

    # TABLE SESSIONS
    ("table_sessions", "id", "BIGINT (PK)", "Primary key, auto-increment", "Unique identifier"),
    ("table_sessions", "table_id", "BIGINT (FK)", "FK to tables.id", "CASCADE; sessions removed if table is deleted"),
    ("table_sessions", "merge_group_id", "BIGINT (FK)", "FK to table_merge_groups.id, nullable", "SET_NULL; session survives merge group dissolution"),
    ("table_sessions", "opened_by_id", "BIGINT (FK)", "FK to users.id, nullable", "SET_NULL; staff who opened the session"),
    ("table_sessions", "customer_account_id", "BIGINT (FK)", "FK to users.id, nullable", "SET_NULL; customer linked to this session"),
    ("table_sessions", "scan_request_id", "BIGINT (FK)", "FK to table_scan_requests.id, nullable", "SET_NULL; QR scan that initiated this session"),
    ("table_sessions", "source", "VARCHAR(20)", "Source: qr/walk_in/waiter/manual", "How the session was started"),
    ("table_sessions", "party_size", "INTEGER", "Number of guests", "Customer count for this seating"),
    ("table_sessions", "guest_label", "VARCHAR(120)", "Readable guest identifier, nullable", "Friendly name for walk-in guests"),
    ("table_sessions", "started_at", "DATETIME", "Session start time", "Auto-set when session is created"),
    ("table_sessions", "ended_at", "DATETIME", "Session end time, nullable", "Set when the bill is paid and table is released"),
    ("table_sessions", "is_active", "BOOLEAN", "Session active flag", "True while customers are seated; false after checkout"),
    ("table_sessions", "created_at", "DATETIME", "Record creation timestamp", "Inherited from BaseModel"),
    ("table_sessions", "updated_at", "DATETIME", "Last modification timestamp", "Inherited from BaseModel"),

    # TABLE SCAN REQUESTS
    ("table_scan_requests", "id", "BIGINT (PK)", "Primary key, auto-increment", "Unique identifier"),
    ("table_scan_requests", "table_id", "BIGINT (FK)", "FK to tables.id", "CASCADE; scan requests removed with table"),
    ("table_scan_requests", "requested_by_id", "BIGINT (FK)", "FK to users.id, nullable", "SET_NULL; identifies the customer who scanned"),
    ("table_scan_requests", "requested_device_id", "VARCHAR(128)", "Device that scanned, nullable", "Guest device ID for session continuity"),
    ("table_scan_requests", "status", "VARCHAR(20)", "Status: pending/approved/blocked", "Moderation state of the scan request"),
    ("table_scan_requests", "note", "VARCHAR(255)", "Staff note, nullable", "Optional staff comment on the request"),
    ("table_scan_requests", "blocked_reason", "VARCHAR(255)", "Block reason, nullable", "Why the scan was rejected (if blocked)"),
    ("table_scan_requests", "created_at", "DATETIME", "Record creation timestamp", "Inherited from BaseModel"),
    ("table_scan_requests", "updated_at", "DATETIME", "Last modification timestamp", "Inherited from BaseModel"),

    # STAFF PAGE REQUESTS
    ("staff_page_requests", "id", "BIGINT (PK)", "Primary key, auto-increment", "Unique identifier"),
    ("staff_page_requests", "table_id", "BIGINT (FK)", "FK to tables.id, nullable", "SET_NULL; request survives table deletion"),
    ("staff_page_requests", "session_id", "BIGINT (FK)", "FK to table_sessions.id, nullable", "SET_NULL; request survives session closure"),
    ("staff_page_requests", "requested_by_id", "BIGINT (FK)", "FK to users.id, nullable", "SET_NULL; customer who requested assistance"),
    ("staff_page_requests", "resolved_by_id", "BIGINT (FK)", "FK to users.id, nullable", "SET_NULL; staff who responded"),
    ("staff_page_requests", "reason", "VARCHAR(20)", "Reason: cleanup/payment/accident", "Categorizes the type of assistance needed"),
    ("staff_page_requests", "status", "VARCHAR(20)", "Status: pending/finished", "Request lifecycle state"),
    ("staff_page_requests", "created_at", "DATETIME", "Record creation timestamp", "Inherited from BaseModel"),
    ("staff_page_requests", "updated_at", "DATETIME", "Last modification timestamp", "Inherited from BaseModel"),

    # NOTIFICATIONS
    ("notifications", "id", "BIGINT (PK)", "Primary key, auto-increment", "Unique identifier"),
    ("notifications", "recipient_id", "BIGINT (FK)", "FK to users.id, nullable", "CASCADE; notification removed if user is deleted"),
    ("notifications", "role_target", "VARCHAR(20)", "Target role, nullable", "Broadcast to all users of this role (waiter/kitchen/etc.)"),
    ("notifications", "title", "VARCHAR(150)", "Notification title", "Brief alert header shown to the user"),
    ("notifications", "message", "TEXT", "Notification body text", "Full message content"),
    ("notifications", "category", "VARCHAR(50)", "Category tag: general/order/table/inventory/etc.", "Used for client-side grouping and filtering"),
    ("notifications", "metadata", "JSON", "Arbitrary JSON payload", "Extra data for the notification action (e.g. order_id, table_id)"),
    ("notifications", "is_read", "BOOLEAN", "Read/unread flag", "Tracks whether the user has seen the notification"),
    ("notifications", "created_at", "DATETIME", "Record creation timestamp", "Inherited from BaseModel"),
    ("notifications", "updated_at", "DATETIME", "Last modification timestamp", "Inherited from BaseModel"),

    # DEBOUNCE RECORDS
    ("debounce_records", "id", "BIGINT (PK)", "Primary key, auto-increment", "Unique identifier"),
    ("debounce_records", "actor_key", "VARCHAR(120)", "Actor identifier (user ID, device ID, etc.)", "Part of the deduplication key; identifies who performed the action"),
    ("debounce_records", "action", "VARCHAR(80)", "Action name", "Part of deduplication key; identifies the action type"),
    ("debounce_records", "object_key", "VARCHAR(120)", "Object identifier, nullable", "Part of deduplication key; identifies what was acted upon"),
    ("debounce_records", "created_at", "DATETIME", "Record creation timestamp", "Inherited from BaseModel"),
    ("debounce_records", "updated_at", "DATETIME", "Last modification timestamp", "Inherited from BaseModel"),

    # AUDIT LOGS
    ("audit_logs", "id", "BIGINT (PK)", "Primary key, auto-increment", "Unique identifier"),
    ("audit_logs", "actor_id", "BIGINT (FK)", "FK to users.id, nullable", "SET_NULL; log survives user deletion"),
    ("audit_logs", "actor_role", "VARCHAR(20)", "Role at time of action, nullable", "Snapshot of the actor's role for historical context"),
    ("audit_logs", "action", "VARCHAR(120)", "Action name/description", "What was done (e.g. 'order.cancelled', 'menu.price_updated')"),
    ("audit_logs", "target_type", "VARCHAR(120)", "Target model/entity name, nullable", "Polymorphic FK target type"),
    ("audit_logs", "target_id", "VARCHAR(64)", "Target record ID, nullable", "Polymorphic FK target identifier"),
    ("audit_logs", "target_label", "VARCHAR(255)", "Human-readable target description, nullable", "Display-friendly target reference"),
    ("audit_logs", "metadata", "JSON", "Arbitrary JSON payload", "Flexible storage for action-specific details"),
    ("audit_logs", "created_at", "DATETIME", "Record creation timestamp", "Inherited from BaseModel"),
    ("audit_logs", "updated_at", "DATETIME", "Last modification timestamp", "Inherited from BaseModel"),

    # GENERATED REPORTS
    ("generated_reports", "id", "BIGINT (PK)", "Primary key, auto-increment", "Unique identifier"),
    ("generated_reports", "generated_by_id", "BIGINT (FK)", "FK to users.id, nullable", "SET_NULL; report survives user deletion"),
    ("generated_reports", "range_type", "VARCHAR(20)", "Range: daily/weekly/monthly/annual/custom", "Defines the report period type"),
    ("generated_reports", "start_at", "DATETIME", "Report range start", "Beginning of the data range for this report"),
    ("generated_reports", "end_at", "DATETIME", "Report range end", "End of the data range for this report"),
    ("generated_reports", "payload", "JSON", "Report data payload", "Stores the generated report results as JSON"),
    ("generated_reports", "created_at", "DATETIME", "Record creation timestamp", "Inherited from BaseModel"),
    ("generated_reports", "updated_at", "DATETIME", "Last modification timestamp", "Inherited from BaseModel"),

    # COST SIMULATIONS
    ("cost_simulations", "id", "BIGINT (PK)", "Primary key, auto-increment", "Unique identifier"),
    ("cost_simulations", "created_by_id", "BIGINT (FK)", "FK to users.id, nullable", "SET_NULL; simulation survives user deletion"),
    ("cost_simulations", "assumptions", "JSON", "Simulation input assumptions", "What-if parameters: ingredient prices, quantities, etc."),
    ("cost_simulations", "results", "JSON", "Simulation output results", "Computed costs and margins based on assumptions"),
    ("cost_simulations", "created_at", "DATETIME", "Record creation timestamp", "Inherited from BaseModel"),
    ("cost_simulations", "updated_at", "DATETIME", "Last modification timestamp", "Inherited from BaseModel"),

    # EXTERNAL SYSTEMS
    ("external_systems", "id", "BIGINT (PK)", "Primary key, auto-increment", "Unique identifier"),
    ("external_systems", "name", "VARCHAR(120)", "System name, unique", "Display name for the integration partner"),
    ("external_systems", "system_type", "VARCHAR(20)", "Type: pos/mobile/kiosk", "Categorizes the integration for routing and handling"),
    ("external_systems", "is_active", "BOOLEAN", "Integration active flag", "False pauses sync without removing configuration"),
    ("external_systems", "created_at", "DATETIME", "Record creation timestamp", "Inherited from BaseModel"),
    ("external_systems", "updated_at", "DATETIME", "Last modification timestamp", "Inherited from BaseModel"),

    # SYNC EVENTS
    ("sync_events", "id", "BIGINT (PK)", "Primary key, auto-increment", "Unique identifier"),
    ("sync_events", "source_system_id", "BIGINT (FK)", "FK to external_systems.id, nullable", "SET_NULL; event survives system configuration removal"),
    ("sync_events", "event_type", "VARCHAR(80)", "Event type name", "Identifies what happened (e.g. 'order.created', 'menu.updated')"),
    ("sync_events", "aggregate_type", "VARCHAR(80)", "Aggregate/model type", "The domain entity affected by this event"),
    ("sync_events", "aggregate_id", "VARCHAR(64)", "Aggregate record ID", "The specific record that changed"),
    ("sync_events", "payload", "JSON", "Event data payload", "Full snapshot or delta of the changed data"),
    ("sync_events", "idempotency_key", "VARCHAR(120)", "Unique idempotency key, nullable", "Prevents duplicate processing of the same event"),
    ("sync_events", "created_at", "DATETIME", "Record creation timestamp", "Inherited from BaseModel"),
    ("sync_events", "updated_at", "DATETIME", "Last modification timestamp", "Inherited from BaseModel"),

    # REALTIME EVENTS
    ("realtime_events", "id", "BIGINT (PK)", "Primary key, auto-increment", "Unique identifier"),
    ("realtime_events", "event_type", "VARCHAR(120)", "Event type name", "Identifies the realtime event for client-side handlers"),
    ("realtime_events", "payload", "JSON", "Event data payload", "Data sent to connected clients"),
    ("realtime_events", "role_target", "VARCHAR(20)", "Target role filter, nullable", "Restricts event to clients with this role"),
    ("realtime_events", "user_id", "BIGINT (FK)", "FK to users.id, nullable", "CASCADE; restricts event to a specific user"),
    ("realtime_events", "created_at", "DATETIME", "Record creation timestamp", "Inherited from BaseModel"),
    ("realtime_events", "updated_at", "DATETIME", "Last modification timestamp", "Inherited from BaseModel"),
]

# ---- Table descriptions ----
table_descriptions = {
    "users": "Central actor table. Stores all account types: customers, waiters, kitchen staff, bar staff, cashiers, and admins. Supports guest flows with is_guest and guest_expires_at. Uses soft-delete (is_deleted) instead of row removal to preserve FK integrity across all related tables.",
    "categories": "Menu category groupings (e.g. Appetizers, Main Course, Desserts, Drinks). Controls the ordering and hierarchy of menu items displayed to customers.",
    "menu_items": "Individual sellable items on the menu. Each item is assigned to a preparation station (kitchen/bar) which determines where the order is routed. is_available is toggled automatically when ingredient stock runs low.",
    "menu_item_ingredients": "Junction table linking menu items to their required ingredients. The quantity_required payload drives automated inventory deduction when an order is placed.",
    "ingredients": "Raw materials tracked in inventory. Each ingredient has a unit of measure (g/ml/pc) and a reorder_level threshold for low-stock alerts.",
    "inventory_batches": "Track individual deliveries/restocks of an ingredient. Enables FIFO cost tracking via quantity_remaining and unit_cost per batch. Expiration dates support spoilage monitoring.",
    "inventory_movements": "Immutable log of every inventory transaction: restocks, deductions, adjustments, and deletions. Links to the actor who performed it and optionally to the order that caused the deduction.",
    "orders": "Core transaction record. Holds the full lifecycle of a customer order with three parallel status tracks (kitchen, bar, cashier) that advance independently. Denormalized fields (placed_by_name, placed_for_name) preserve context even if related users are deleted.",
    "order_items": "Line items within an order. Denormalizes item_name, station, and unit_price at time of order so historical receipts remain accurate after menu changes.",
    "order_status_logs": "Immutable audit trail for order status transitions. Each log records who changed the status, what it changed to, and optional metadata for full traceability.",
    "feedback_entries": "Customer feedback and problem reports. Supports two entry types (feedback/report) with a status lifecycle (open/reviewed/resolved) for staff follow-up.",
    "order_playlists": "Named saved-order templates owned by a user. Lets customers save favorite combinations for quick reordering. Unique on (owner_id, name).",
    "order_playlist_items": "Junction table linking a playlist to menu items with a default quantity. Unique on (playlist_id, menu_item_id) prevents duplicate entries.",
    "tables": "Physical tables in the venue. Each has a unique display identifier, seating capacity, status for floorplan management, zone for area grouping, and an auto-generated QR code for customer scanning.",
    "table_merge_groups": "Groups of tables that have been physically merged to accommodate larger parties. The M:N relationship with tables (via table_merge_group_tables) allows flexible combinations.",
    "table_merge_group_tables": "Implicit M:N junction table (auto-generated by Django) linking merge groups to tables. No payload — the relationship itself carries no extra data.",
    "table_sessions": "A seating session representing a customer occupation of a table. Tracks who opened it, which customer is seated, how they arrived (source), party size, and active duration.",
    "table_scan_requests": "QR code scan intents from customers. Each request records which table was scanned, by whom, and the moderation outcome (pending/approved/blocked).",
    "staff_page_requests": "Customer requests for staff assistance. Includes the reason (cleanup/payment/accident) and tracks resolution by staff. Links to both table and session for context.",
    "notifications": "In-app notification records. Can target a specific user via recipient_id or broadcast to all users of a given role via role_target. Category field enables client-side grouping.",
    "debounce_records": "Idempotency guard table. Prevents duplicate processing of the same action by the same actor on the same object. The unique constraint on (actor_key, action, object_key) serves as the lock. No FKs — uses loose strings for speed and flexibility.",
    "audit_logs": "Generic action audit trail. Uses a polymorphic FK pattern (target_type + target_id + target_label) to reference any entity without formal FK constraints. Captures actor, action, role snapshot, and metadata.",
    "generated_reports": "Stores report generation requests and their results. Supports daily, weekly, monthly, annual, and custom date ranges. Report data is stored as JSON in the payload field.",
    "cost_simulations": "What-if cost analysis scenarios. Users input assumption parameters (ingredient prices, quantities, waste %) as JSON, and the system computes projected costs and margins — stored as results JSON.",
    "external_systems": "Registered external integrations (POS systems, mobile apps, kiosks). Each has a type label and active flag to enable/disable sync without removing configuration.",
    "sync_events": "Outbox-pattern event log for external system synchronization. Events are written during the main request and consumed asynchronously by a background worker. The idempotency_key prevents duplicate processing.",
    "realtime_events": "Server-Sent Events (SSE) data store. Events are created by services (orders, inventory, tables) and consumed by connected clients via polling-based SSE at /api/realtime/stream/. Supports targeting by role or specific user.",
}

# ---- Build document ----
current_table = None

for table, field, dtype, desc, expl in entries:
    if table != current_table:
        current_table = table
        doc.add_heading(table, level=1)
        doc.add_paragraph(table_descriptions.get(table, ""))
        table_obj = doc.add_table(rows=1, cols=3)
        table_obj.style = "Light Grid Accent 1"
        table_obj.autofit = False
        table_obj.alignment = WD_TABLE_ALIGNMENT.CENTER
        for cell in table_obj.rows[0].cells:
            cell.width = Inches(2.15)

        # Header row
        hdr = table_obj.rows[0].cells
        headers = ["Field", "Data Type", "Description"]
        for i, text in enumerate(headers):
            hdr[i].text = text
            for paragraph in hdr[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

    row = table_obj.add_row().cells
    row[0].text = field
    row[1].text = dtype
    row[2].text = desc

# ---- Save ----
doc.save("MicroBlend_Data_Dictionary.docx")
print("Done: MicroBlend_Data_Dictionary.docx")
